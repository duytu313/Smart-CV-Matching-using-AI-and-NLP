"""
CV/Resume Parsing Service.
Handles extraction of text and structured data from PDF, DOCX, and image files.
"""

import re
import io
import sys
import os
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from .skill_service import SkillService

# Thêm đường dẫn để import config
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
try:
    from config import settings
except ImportError:
    # Fallback settings nếu chưa có config
    class DefaultSettings:
        TESSERACT_PATH = ""
        OCR_LANGUAGE = "eng"
        TESSERACT_CONFIG = "--oem 3 --psm 6"
        OCR_IMAGE_MIN_WIDTH = 1000
        OCR_CONTRAST_ENHANCE_FACTOR = 2.0
    
    settings = DefaultSettings()


@dataclass
class ParsedResume:
    """Structured resume data."""
    raw_text: str
    skills: List[str]
    experience_years: Optional[float]
    education: List[Dict]
    contact_info: Dict
    sections: Dict[str, str]


class CVParser:
    """
    Service for parsing and extracting information from resumes.
    Supports PDF, DOCX, and image formats (PNG, JPG, JPEG) with OCR.
    """
    
    # Patterns for information extraction
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_PATTERN = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
    LINKEDIN_PATTERN = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9\-_]+)'
    GITHUB_PATTERN = r'(?:github\.com/|github:?\s*)([a-zA-Z0-9\-_]+)'
    
    # Section headers to identify
    SECTION_HEADERS = [
        'experience', 'work experience', 'professional experience', 'employment',
        'education', 'academic background', 'qualifications',
        'skills', 'technical skills', 'core competencies', 'expertise',
        'projects', 'personal projects', 'key projects',
        'certifications', 'certificates', 'licenses',
        'summary', 'professional summary', 'objective', 'profile',
        'publications', 'awards', 'achievements', 'honors',
        'languages', 'interests', 'hobbies', 'references'
    ]
    
    # Education keywords
    DEGREES = [
        'phd', 'ph.d', 'doctorate', 'doctoral',
        'master', 'masters', 'mba', 'msc', 'ma', 'ms', 'm.s.',
        'bachelor', 'bachelors', 'bsc', 'ba', 'bs', 'b.s.', 'btech', 'be',
        'associate', 'diploma', 'certificate'
    ]
    
    @classmethod
    def parse_pdf(cls, file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return ""
    
    @classmethod
    def parse_docx(cls, file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return ""
    
    @classmethod
    def parse_image(cls, file_content: bytes) -> str:
        """
        Extract text from image file using OCR (Tesseract).
        
        Args:
            file_content: Image file bytes
            
        Returns:
            Extracted text from image
        """
        try:
            from PIL import Image
            import pytesseract
            
            # Set Tesseract path if configured
            if settings.TESSERACT_PATH:
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
            
            # Open image from bytes
            image = Image.open(io.BytesIO(file_content))
            
            # Preprocess image for better OCR results
            # Convert to RGB if necessary
            if image.mode not in ('L', 'RGB'):
                image = image.convert('RGB')
            
            # Use pytesseract to extract text with language setting
            custom_config = f"--oem 3 --psm 6 -l {settings.OCR_LANGUAGE}"
            extracted_text = pytesseract.image_to_string(image, config=custom_config)
            
            return extracted_text.strip()
        except ImportError as e:
            print(f"OCR libraries not installed: {e}")
            print("Please install: pip install Pillow pytesseract")
            return ""
        except Exception as e:
            print(f"Error parsing image with OCR: {e}")
            return ""
    
    @classmethod
    def parse_image_with_preprocessing(cls, file_content: bytes) -> str:
        """
        Extract text from image with advanced preprocessing for better accuracy.
        
        Args:
            file_content: Image file bytes
            
        Returns:
            Extracted text from image
        """
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            import pytesseract
            
            # Set Tesseract path if configured
            if settings.TESSERACT_PATH:
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
            
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast using config value
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(settings.OCR_CONTRAST_ENHANCE_FACTOR)
            
            # Apply sharpening filter
            image = image.filter(ImageFilter.SHARPEN)
            
            # Resize image if too small (improves OCR accuracy)
            if image.width < settings.OCR_IMAGE_MIN_WIDTH:
                ratio = settings.OCR_IMAGE_MIN_WIDTH / image.width
                new_size = (int(image.width * ratio), int(image.height * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Extract text with custom configuration including language
            custom_config = f"--oem 3 --psm 6 -l {settings.OCR_LANGUAGE} -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789@.,:;/-_() "
            extracted_text = pytesseract.image_to_string(image, config=custom_config)
            
            return extracted_text.strip()
        except Exception as e:
            print(f"Error in advanced image parsing: {e}")
            return cls.parse_image(file_content)  # Fallback to basic OCR
    
    @classmethod
    def parse_file(cls, file_content: bytes, filename: str) -> str:
        """
        Parse a resume file based on its extension.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            
        Returns:
            Extracted text
        """
        filename_lower = filename.lower()
        
        # Image formats
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.webp']
        
        if filename_lower.endswith('.pdf'):
            return cls.parse_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return cls.parse_docx(file_content)
        elif filename_lower.endswith('.doc'):
            # For .doc files, we'd need additional libraries
            raise ValueError("Legacy .doc format not supported. Please use .docx or .pdf")
        elif filename_lower.endswith('.txt'):
            return file_content.decode('utf-8', errors='ignore')
        elif any(filename_lower.endswith(ext) for ext in image_extensions):
            # Try advanced OCR first, fallback to basic OCR
            text = cls.parse_image_with_preprocessing(file_content)
            if not text:
                text = cls.parse_image(file_content)
            if not text:
                raise ValueError("Could not extract text from image. Please ensure the image is clear and contains text.")
            return text
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    
    @classmethod
    def extract_contact_info(cls, text: str) -> Dict:
        """
        Extract contact information from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with contact details
        """
        contact = {}
        
        # Extract email
        emails = re.findall(cls.EMAIL_PATTERN, text)
        if emails:
            contact['email'] = emails[0]
        
        # Extract phone
        phones = re.findall(cls.PHONE_PATTERN, text)
        if phones:
            contact['phone'] = phones[0]
        
        # Extract LinkedIn
        linkedin = re.search(cls.LINKEDIN_PATTERN, text, re.IGNORECASE)
        if linkedin:
            contact['linkedin'] = linkedin.group(1)
        
        # Extract GitHub
        github = re.search(cls.GITHUB_PATTERN, text, re.IGNORECASE)
        if github:
            contact['github'] = github.group(1)
        
        return contact
    
    @classmethod
    def extract_experience_years(cls, text: str) -> Optional[float]:
        """
        Estimate years of experience from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Estimated years of experience
        """
        text_lower = text.lower()
        
        # Look for explicit statements
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience\s*:\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*(?:of\s*)?professional',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return float(match.group(1))
        
        # Try to count years from date ranges
        date_ranges = re.findall(
            r'(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|present|current)',
            text_lower
        )
        
        if date_ranges:
            total_years = 0
            for start, end in date_ranges:
                try:
                    start_year = int(start)
                    if end in ['present', 'current']:
                        end_year = 2024
                    else:
                        end_year = int(end)
                    total_years += max(0, end_year - start_year)
                except:
                    pass
            
            if total_years > 0:
                return float(total_years)
        
        return None
    
    @classmethod
    def extract_education(cls, text: str) -> List[Dict]:
        """
        Extract education information from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of education entries
        """
        education = []
        text_lower = text.lower()
        
        # Find education section
        edu_start = -1
        for header in ['education', 'academic', 'qualifications']:
            idx = text_lower.find(header)
            if idx != -1:
                edu_start = idx
                break
        
        if edu_start == -1:
            # Try to find degrees anywhere in the text
            for degree in cls.DEGREES:
                if degree in text_lower:
                    # Found a degree mention
                    lines = text.split('\n')
                    for line in lines:
                        if degree in line.lower():
                            education.append({
                                'degree': line.strip(),
                                'institution': '',
                                'year': ''
                            })
        else:
            # Extract from education section
            section_text = text[edu_start:edu_start + 1000]  # Get next 1000 chars
            lines = section_text.split('\n')
            
            current_edu = {}
            for line in lines[1:]:  # Skip header
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a new section
                if any(header in line.lower() for header in cls.SECTION_HEADERS if header not in ['education', 'academic']):
                    break
                
                # Look for degree
                has_degree = any(deg in line.lower() for deg in cls.DEGREES)
                if has_degree:
                    if current_edu:
                        education.append(current_edu)
                    current_edu = {'degree': line, 'institution': '', 'year': ''}
                elif current_edu:
                    # This might be the institution or year
                    year_match = re.search(r'20\d{2}|19\d{2}', line)
                    if year_match:
                        current_edu['year'] = year_match.group()
                    else:
                        current_edu['institution'] = line
            
            if current_edu:
                education.append(current_edu)
        
        return education
    
    @classmethod
    def extract_sections(cls, text: str) -> Dict[str, str]:
        """
        Split resume into sections based on headers.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary mapping section names to content
        """
        sections = {}
        lines = text.split('\n')
        
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            is_header = False
            for header in cls.SECTION_HEADERS:
                if line_lower == header or line_lower.startswith(header + ':'):
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = header.split()[0]  # Use first word as key
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Add the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    @classmethod
    def parse_resume(cls, file_content: bytes, filename: str) -> ParsedResume:
        """
        Parse a resume file and extract all structured information.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            
        Returns:
            ParsedResume object with all extracted data
        """
        # Extract raw text
        raw_text = cls.parse_file(file_content, filename)
        
        if not raw_text:
            raise ValueError("Could not extract text from the file")
        
        # Extract skills
        skills = SkillService.extract_skills(raw_text)
        
        # Extract other information
        experience_years = cls.extract_experience_years(raw_text)
        education = cls.extract_education(raw_text)
        contact_info = cls.extract_contact_info(raw_text)
        sections = cls.extract_sections(raw_text)
        
        return ParsedResume(
            raw_text=raw_text,
            skills=skills,
            experience_years=experience_years,
            education=education,
            contact_info=contact_info,
            sections=sections
        )